# judge_docx_generator.py — генерация «Карточки учёта судейской деятельности
# спортивного судьи» в Word по образцу из data/
import copy
import os
from io import BytesIO

import docx
from docx.oxml.ns import qn
from docx.table import _Cell

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, "data", "Учетная карточка судьи СС1К Серганова СА.docx")

CATEGORY_ABBREVIATIONS = {
    "Спортивный судья всероссийской категории": "ССВК",
    "Спортивный судья первой категории": "СС1К",
    "Спортивный судья второй категории": "СС2К",
    "Спортивный судья третьей категории": "СС3К",
}


def _tcs(tr):
    return tr.findall(qn("w:tc"))


def _set_text(tr, tc_idx, table, value):
    tc = _tcs(tr)[tc_idx]
    cell = _Cell(tc, table)
    value = "" if value is None else str(value)
    lines = value.split("\n")

    paragraph = cell.paragraphs[0]
    for extra_p in cell.paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)

    runs = paragraph.runs
    if runs:
        run = runs[0]
        for extra_r in runs[1:]:
            extra_r._element.getparent().remove(extra_r._element)
    else:
        run = paragraph.add_run()

    run.text = lines[0]
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)


def _remove_grid_column(table, col_index):
    """Удаляет col_index-ю колонку сетки таблицы целиком: одну gridCol и
    соответствующую ячейку (или -1 к gridSpan объединённой ячейки) в каждой
    физической строке таблицы."""
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    grid_cols = tblGrid.findall(qn("w:gridCol"))
    tblGrid.remove(grid_cols[col_index])

    for tr in tbl.findall(qn("w:tr")):
        pos = 0
        for tc in tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            gridSpan = tcPr.find(qn("w:gridSpan")) if tcPr is not None else None
            span = int(gridSpan.get(qn("w:val"))) if gridSpan is not None else 1
            if pos <= col_index < pos + span:
                if span > 1:
                    gridSpan.set(qn("w:val"), str(span - 1))
                else:
                    tr.remove(tc)
                break
            pos += span


def _replace_repeating_rows(table, header_row_count, group_size, records, fill_group):
    """Заменяет строки-образцы таблицы (после header_row_count строк заголовка)
    на group_size-строчные блоки — по одному на каждую запись из records."""
    tbl = table._tbl
    all_trs = tbl.findall(qn("w:tr"))
    header_trs = all_trs[:header_row_count]
    sample_trs = all_trs[header_row_count:]
    template_group = sample_trs[:group_size]

    for tr in sample_trs:
        tbl.remove(tr)

    anchor = header_trs[-1]
    for record in records:
        new_group = [copy.deepcopy(tr) for tr in template_group]
        for tr in new_group:
            anchor.addnext(tr)
            anchor = tr
        fill_group(new_group, record)


def _fmt_date(value):
    return value.strftime("%d.%m.%Y") if value else ""


def _date_parts(value):
    if not value:
        return "", "", ""
    return value.strftime("%d"), value.strftime("%m"), value.strftime("%Y")


def generate_judge_card(judge, settings):
    document = docx.Document(TEMPLATE_PATH)
    table0, table1, table2 = document.tables[0], document.tables[1], document.tables[2]

    # Убираем столбец «Дата» (участника), стоящий под группой «Участие в
    # теоретической подготовке в качестве Лектора»
    _remove_grid_column(table1, 6)

    # --- Шапка: вид спорта, организация ---
    _set_text(table0.rows[0]._tr, 2, table0, settings.sport_name if settings else "")
    _set_text(table0.rows[1]._tr, 2, table0, settings.sport_code if settings else "")
    _set_text(table0.rows[11]._tr, 1, table0, settings.organization_name if settings else "")
    _set_text(table0.rows[11]._tr, 3, table0, settings.organization_address if settings else "")
    _set_text(table0.rows[11]._tr, 5, table0, settings.organization_contacts if settings else "")

    # --- Личные данные судьи ---
    _set_text(table0.rows[2]._tr, 1, table0, judge.last_name)
    _set_text(table0.rows[2]._tr, 3, table0, judge.first_name)
    _set_text(table0.rows[2]._tr, 5, table0, judge.middle_name)

    birth_day, birth_month, birth_year = _date_parts(judge.birth_date)
    _set_text(table0.rows[4]._tr, 6, table0, birth_day)
    _set_text(table0.rows[4]._tr, 7, table0, birth_month)
    _set_text(table0.rows[4]._tr, 8, table0, birth_year)

    _set_text(table0.rows[4]._tr, 1, table0, judge.region)
    _set_text(table0.rows[4]._tr, 3, table0, judge.municipality)
    _set_text(table0.rows[4]._tr, 5, table0, judge.rank)

    start_day, start_month, start_year = _date_parts(judge.judging_start_date)
    _set_text(table0.rows[7]._tr, 2, table0, start_day)
    _set_text(table0.rows[7]._tr, 3, table0, start_month)
    _set_text(table0.rows[7]._tr, 4, table0, start_year)

    _set_text(table0.rows[7]._tr, 1, table0, judge.education)
    _set_text(table0.rows[8]._tr, 1, table0, judge.workplace)
    _set_text(table0.rows[9]._tr, 1, table0, judge.contacts)

    # --- История квалификационных категорий (table0, 1 строка на запись) ---
    def fill_category_group(group_trs, record):
        tr = group_trs[0]
        _set_text(tr, 0, table0, record.category)
        _set_text(tr, 1, table0, record.action)
        _set_text(tr, 2, table0, _fmt_date(record.date))
        _set_text(tr, 3, table0, record.document_number)
        _set_text(tr, 4, table0, record.issuing_organization)
        _set_text(tr, 5, table0, record.signed_by)
        _set_text(tr, 6, table0, record.record_keeper)

    _replace_repeating_rows(table0, 14, 1, judge.category_records, fill_category_group)

    # --- Теоретическая подготовка, сдача квалификационного зачёта (table1, 3 строки на запись) ---
    # Столбец «Дата (участник)» удалён из таблицы (_remove_grid_column выше),
    # поэтому индексы столбцов с лектора и дальше сдвинуты на 1 влево.
    def fill_training_group(group_trs, record):
        tr0, tr2 = group_trs[0], group_trs[2]
        _set_text(tr0, 0, table1, record.seminar_name)
        _set_text(tr0, 1, table1, record.seminar_date)
        _set_text(tr0, 2, table1, record.organizer)
        _set_text(tr0, 3, table1, record.location)
        _set_text(tr0, 4, table1, record.participant_category)
        _set_text(tr0, 5, table1, record.participant_score)
        _set_text(tr0, 6, table1, record.lecturer_category)
        _set_text(tr0, 7, table1, record.lecturer_score)
        _set_text(tr0, 8, table1, record.lecturer_date)
        _set_text(tr0, 9, table1, record.exam_protocol_number)
        _set_text(tr0, 10, table1, record.exam_score)
        _set_text(tr0, 11, table1, record.record_date)
        _set_text(tr2, 11, table1, record.record_keeper)

    _replace_repeating_rows(table1, 2, 3, judge.training_records, fill_training_group)

    # --- Практика судейства официальных спортивных соревнований (table2, 2 строки на запись) ---
    def fill_competition_group(group_trs, record):
        tr0, tr1 = group_trs[0], group_trs[1]
        _set_text(tr0, 0, table2, record.event_date)
        _set_text(tr0, 1, table2, record.location)
        _set_text(tr0, 2, table2, record.judge_position)
        _set_text(tr0, 3, table2, record.competition_name)
        _set_text(tr0, 4, table2, record.score)
        _set_text(tr0, 5, table2, record.record_date)
        _set_text(tr1, 3, table2, record.competition_status)
        _set_text(tr1, 5, table2, record.record_keeper)

    _replace_repeating_rows(table2, 1, 2, judge.competition_records, fill_competition_group)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def judge_card_filename(judge):
    abbr = CATEGORY_ABBREVIATIONS.get(judge.current_category, "")
    initials = "".join(f"{p[0]}." for p in (judge.first_name, judge.middle_name) if p)
    parts = ["Карточка", "судьи", abbr, judge.last_name, initials]
    return " ".join(p for p in parts if p).strip() + ".docx"
