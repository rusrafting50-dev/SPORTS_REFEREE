# judge_docx_importer.py — импорт данных из готовой docx-карточки судьи
# («Карточка учёта судейской деятельности спортивного судьи») в новую
# запись Judge с историей категорий/подготовки/практики
import re
from datetime import date

import docx
from docx.oxml.ns import qn

from models import Judge, JudgeCategoryRecord, JudgeCompetitionRecord, JudgeTrainingRecord

_DATE_RE = re.compile(r"(\d{1,2})\.(\d{1,2})\.(\d{4})")


def _cell(table, row, col):
    """Текст ячейки по индексу физической <w:tc> в строке (а не индексу
    сетки таблицы) — так же, как ячейки заполняются в judge_docx_generator.py,
    поскольку строки шапки/истории категорий содержат объединённые по
    горизонтали ячейки (gridSpan), из-за чего индекс сетки не совпадает
    с порядковым номером ячейки."""
    try:
        tr = table.rows[row]._tr
    except IndexError:
        return ""
    tcs = tr.findall(qn("w:tc"))
    if col >= len(tcs):
        return ""
    texts = tcs[col].findall(".//" + qn("w:t"))
    return "".join(t.text or "" for t in texts).strip()


def _row_blank(table, row, cols):
    return all(not _cell(table, row, c) for c in cols)


def _parse_date_parts(day, month, year):
    if not (day.isdigit() and month.isdigit() and year.isdigit()):
        return None
    try:
        return date(int(year), int(month), int(day))
    except ValueError:
        return None


def _parse_ddmmyyyy(text):
    match = _DATE_RE.search(text or "")
    if not match:
        return None
    day, month, year = match.groups()
    try:
        return date(int(year), int(month), int(day))
    except ValueError:
        return None


def parse_judge_card(file_obj):
    """Разбирает docx-карточку учёта судейской деятельности и возвращает
    новый объект Judge со связанными записями истории (без сохранения в БД)."""
    document = docx.Document(file_obj)
    table0, table1, table2 = document.tables[0], document.tables[1], document.tables[2]

    judge = Judge(
        last_name=_cell(table0, 2, 1),
        first_name=_cell(table0, 2, 3),
        middle_name=_cell(table0, 2, 5) or None,
        region=_cell(table0, 4, 1) or None,
        municipality=_cell(table0, 4, 3) or None,
        rank=_cell(table0, 4, 5) or None,
        education=_cell(table0, 7, 1) or None,
        workplace=_cell(table0, 8, 1) or None,
        contacts=_cell(table0, 9, 1) or None,
        birth_date=_parse_date_parts(_cell(table0, 4, 6), _cell(table0, 4, 7), _cell(table0, 4, 8)),
        judging_start_date=_parse_date_parts(_cell(table0, 7, 2), _cell(table0, 7, 3), _cell(table0, 7, 4)),
    )

    # --- История квалификационных категорий (table0, с 14-й строки, 1 строка на запись) ---
    for row in range(14, len(table0.rows)):
        if _row_blank(table0, row, range(7)):
            continue
        judge.category_records.append(JudgeCategoryRecord(
            category=_cell(table0, row, 0) or None,
            action=_cell(table0, row, 1) or None,
            date=_parse_ddmmyyyy(_cell(table0, row, 2)),
            document_number=_cell(table0, row, 3) or None,
            issuing_organization=_cell(table0, row, 4) or None,
            signed_by=_cell(table0, row, 5) or None,
            record_keeper=_cell(table0, row, 6) or None,
        ))

    # --- Теоретическая подготовка (table1, с 2-й строки, группы по 3 строки) ---
    # В стандартном образце 13 столбцов (со столбцом "Дата" участника), в карточках,
    # выгруженных этим приложением, — 12 (столбец убран из отображения).
    has_participant_date = len(table1.columns) >= 13
    for row in range(2, len(table1.rows), 3):
        if row + 2 >= len(table1.rows):
            break
        if _row_blank(table1, row, range(6)):
            continue
        if has_participant_date:
            judge.training_records.append(JudgeTrainingRecord(
                seminar_name=_cell(table1, row, 0) or None,
                seminar_date=_cell(table1, row, 1) or None,
                organizer=_cell(table1, row, 2) or None,
                location=_cell(table1, row, 3) or None,
                participant_category=_cell(table1, row, 4) or None,
                participant_score=_cell(table1, row, 5) or None,
                participant_date=_cell(table1, row, 6) or None,
                lecturer_category=_cell(table1, row, 7) or None,
                lecturer_score=_cell(table1, row, 8) or None,
                lecturer_date=_cell(table1, row, 9) or None,
                exam_protocol_number=_cell(table1, row, 10) or None,
                exam_score=_cell(table1, row, 11) or None,
                record_date=_cell(table1, row, 12) or None,
                record_keeper=_cell(table1, row + 2, 12) or None,
            ))
        else:
            judge.training_records.append(JudgeTrainingRecord(
                seminar_name=_cell(table1, row, 0) or None,
                seminar_date=_cell(table1, row, 1) or None,
                organizer=_cell(table1, row, 2) or None,
                location=_cell(table1, row, 3) or None,
                participant_category=_cell(table1, row, 4) or None,
                participant_score=_cell(table1, row, 5) or None,
                lecturer_category=_cell(table1, row, 6) or None,
                lecturer_score=_cell(table1, row, 7) or None,
                lecturer_date=_cell(table1, row, 8) or None,
                exam_protocol_number=_cell(table1, row, 9) or None,
                exam_score=_cell(table1, row, 10) or None,
                record_date=_cell(table1, row, 11) or None,
                record_keeper=_cell(table1, row + 2, 11) or None,
            ))

    # --- Практика судейства официальных спортивных соревнований (table2, с 1-й строки, группы по 2 строки) ---
    for row in range(1, len(table2.rows), 2):
        if row + 1 >= len(table2.rows):
            break
        if _row_blank(table2, row, range(3)):
            continue
        judge.competition_records.append(JudgeCompetitionRecord(
            event_date=_cell(table2, row, 0) or None,
            location=_cell(table2, row, 1) or None,
            judge_position=_cell(table2, row, 2) or None,
            competition_name=_cell(table2, row, 3) or None,
            score=_cell(table2, row, 4) or None,
            record_date=_cell(table2, row, 5) or None,
            competition_status=_cell(table2, row + 1, 3) or None,
            record_keeper=_cell(table2, row + 1, 5) or None,
        ))

    return judge
