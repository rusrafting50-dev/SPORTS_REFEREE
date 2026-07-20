# seminar_polozhenie_generator.py — генерация документа «Положение о семинаре»
# Текст положения воспроизводит образец data/Положение семинар СС1К.doc дословно;
# меняются (подставляются из формы «Положение о семинаре») только фрагменты,
# выделенные в образце полужирным шрифтом — даты, наименования, ФИО, суммы и т.п.
from io import BytesIO

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

SEMINAR_CATEGORY_GENITIVE = {
    "Всероссийская": "всероссийской",
    "Первая": "первой",
    "Вторая": "второй",
    "Третья": "третьей",
    "Юный судья": "«Юный судья»",
}

# Категория на ступень ниже присваиваемой — для раздела 4 (кто допускается к семинару)
CATEGORY_PREV_TIER = {
    "Всероссийская": "Первая",
    "Первая": "Вторая",
    "Вторая": "Третья",
    "Третья": None,
    "Юный судья": None,
}

# Значение поля «Количество часов» -> согласованная форма перед словом «программе»
PROGRAM_HOURS_ADJECTIVE = {
    "16-ти часовая": "16-ти часовой",
    "12-ти часовая": "12-ти часовой",
    "10-тичасовая": "10-ти часовой",
}

# Значение поля «Количество часов» -> форма числительного перед «академических часов»
PROGRAM_HOURS_NUMERAL = {
    "16-ти часовая": "16-ти",
    "12-ти часовая": "12-ти",
    "10-тичасовая": "10-ти",
}

RU_MONTHS_GENITIVE = [
    "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]

# Дательный падеж вида спорта для оборотов «по рафтингу» — известен только для рафтинга,
# для прочих видов спорта используется именительный падеж (без склонения).
SPORT_DATIVE = {"рафтинг": "рафтингу"}


def _sport_dative(sport_name):
    return SPORT_DATIVE.get(sport_name, sport_name)


def _fmt_date_ru(value):
    if not value:
        return ""
    return f"{value.day} {RU_MONTHS_GENITIVE[value.month - 1]} {value.year} года"


def _clean(text):
    """Убирает конечную точку из значения поля, чтобы не задваивалась с точкой шаблона."""
    return (text or "").strip().rstrip(".")


def _initials_from_full_name(full_name):
    """«Милевский Евгений Вадимович» -> «Е. В. Милевский» (для подписи в грифе утверждения)."""
    parts = (full_name or "").split()
    if len(parts) < 2:
        return full_name or ""
    last_name, rest = parts[0], parts[1:]
    initials = " ".join(f"{p[0]}." for p in rest if p)
    return f"{initials} {last_name}".strip()


def _org_name_dative(nominative):
    """«Региональная общественная организация «Федерация...»» -> «региональной
    общественной организации «Федерация...»» (родительный/дательный падеж — для
    грифа «УТВЕРЖДАЮ»; само наименование в кавычках не склоняется). Полное
    наименование федерации (раздел 1) при этом остаётся в именительном падеже —
    в этой форме оно используется в разделе «Организаторы»."""
    nominative = (nominative or "").strip()
    if not nominative:
        return ""
    quote_pos = nominative.find("«")
    prefix, quoted = (nominative, "") if quote_pos == -1 else (nominative[:quote_pos], nominative[quote_pos:])
    words = []
    for word in prefix.split():
        if word.endswith("ая"):
            words.append(word[:-2] + "ой")
        elif word.endswith("ия") or word.endswith("ция"):
            words.append(word[:-1] + "и")
        else:
            words.append(word)
    if words:
        words[0] = words[0][0].lower() + words[0][1:]
    return " ".join(words + ([quoted] if quoted else []))


def _region_dative(nominative):
    """«Московская область» -> «Московской области» (дательный падеж — для
    п. 1.1 и п. 4.1 положения)."""
    nominative = (nominative or "").strip()
    if not nominative:
        return ""
    words = []
    for word in nominative.split():
        if word.endswith("ая"):
            words.append(word[:-2] + "ой")
        elif word.endswith("ь"):
            words.append(word[:-1] + "и")
        else:
            words.append(word)
    return " ".join(words)


def _seminar_name_prepositional(name):
    """«Московский областной региональный семинар по подготовке...» -> «Московском
    областном региональном семинаре по подготовке...» — предложный падеж для
    заголовка «ПОЛОЖЕНИЕ о ...». Склоняются только слово «семинар» и определения
    перед ним (обычно название субъекта РФ и уровня семинара); всё, что идёт после
    «семинар» (обороты со своими предлогами — «по подготовке», «по виду спорта» и
    т.п.), не меняется, так как их падеж не зависит от внешнего предлога «о»."""
    name = (name or "").strip()
    if not name:
        return ""
    words = name.split()
    try:
        idx = next(i for i, w in enumerate(words) if w.casefold() == "семинар")
    except StopIteration:
        return name

    declined = []
    for word in words[:idx]:
        if word.endswith(("ий", "ый", "ой")):
            declined.append(word[:-2] + "ом")
        else:
            declined.append(word)
    declined.append(words[idx] + "е")
    declined.extend(words[idx + 1:])
    return " ".join(declined)


def build_polozhenie_data(seminar, lecturers):
    """Собирает содержимое положения о семинаре — единый источник данных и для
    страницы печати в браузере, и для генерации .docx. Текст воспроизводит
    образец data/Положение семинар СС1К.doc дословно; переменные фрагменты
    (в образце — полужирные) подставляются из полей формы «Положение о семинаре»."""
    sport = "рафтинг"
    sport_dat = _sport_dative(sport)
    cat = SEMINAR_CATEGORY_GENITIVE.get(seminar.category, "")

    fed_full = seminar.polozhenie_federation_full_name or ""
    fed_short = seminar.polozhenie_federation_short_name or ""
    fed_region = seminar.polozhenie_federation_region or ""
    fed_region_dative = _region_dative(fed_region)
    period = seminar.polozhenie_period or ""
    location = seminar.polozhenie_location or ""
    hours_value = seminar.polozhenie_program_hours or ""
    hours_adj = PROGRAM_HOURS_ADJECTIVE.get(hours_value, hours_value)
    hours_num = PROGRAM_HOURS_NUMERAL.get(hours_value, hours_value)

    seminar_name = seminar.name or ""
    title_sub = f"о {_seminar_name_prepositional(seminar_name)}" if seminar_name else ""

    # 1. Цели и задачи
    section1 = [
        (
            f"1.1. {seminar_name} (далее по тексту - семинар) "
            f"проводится с целью подготовки и повышения квалификации спортивных судей {cat} категории по {sport_dat}"
            + (f", а также повышения качества проведения официальных спортивных соревнований по {sport_dat}"
               f" в {fed_region_dative}." if fed_region else f", а также повышения качества проведения официальных "
               f"спортивных соревнований по {sport_dat}.")
        ),
        "1.2. Основными задачами семинара являются:",
        (
            f"- выполнение требований по прохождению теоретической подготовки для присвоения (подтверждения) "
            f"квалификационной категории «спортивный судья {cat} категории» в соответствии с квалификационными "
            f"требованиями к спортивным судьям по виду спорта «{sport}»;"
        ),
        (
            f"- выполнение требований по сдаче квалификационного зачета для присвоения (подтверждения) "
            f"квалификационной категории «спортивный судья {cat} категории» в соответствии с квалификационными "
            f"требованиями к спортивным судьям по виду спорта «{sport}»;"
        ),
        (
            f"- обобщение и трансляция передового опыта организации и проведения муниципальных и региональных "
            f"официальных спортивных соревнований по виду спорта «{sport}» для повышения качества организации "
            f"и проведения официальных спортивных соревнований;"
        ),
        (
            f"- совершенствование требований к обеспечению безопасности, постановке дистанций, организации "
            f"судейства и работы секретариата на соревнованиях по виду спорта «{sport}»."
        ),
    ]

    # 2. Время и место проведения
    section2 = []
    if period:
        section2.append(f"2.1. Сроки проведения семинара: {period}.")
    if location:
        section2.append(f"2.2. Место проведения семинара: {location}.")

    # 3. Организаторы
    section3 = []
    if fed_full:
        line = f"3.1. Организатором семинара является {fed_full}"
        if fed_short:
            line += f" (далее по тексту {fed_short})"
        section3.append(line + ".")
    section3.append(
        f"3.2. Непосредственное проведение семинара возлагается на руководителя семинара, утвержденного "
        f"президиумом {fed_short}, и на преподавательский состав семинара, утвержденный руководителем "
        f"семинара. Список преподавательского состава приведен в приложении 1 к настоящему положению."
    )
    leader_bits = ", ".join(
        p for p in (seminar.leader_full_name, seminar.leader_category, seminar.leader_region) if p
    )
    if leader_bits:
        line = f"3.3. Руководитель семинара – {leader_bits}"
        if seminar.leader_phone:
            line += f", тел. {seminar.leader_phone}"
        section3.append(line + ".")

    # 4. Требования к участникам семинара
    section4 = [
        f"4.1. К участию в семинаре допускаются спортивные судьи, учет судейской работы которых "
        f"осуществляется в {fed_short}, в местных спортивных федерациях {fed_region_dative} и в местных "
        f"отделениях {fed_short}, а также других субъектов Российской Федерации.",
        "4.2. К участию в семинаре допускаются:",
    ]
    prev_category = CATEGORY_PREV_TIER.get(seminar.category)
    if prev_category:
        prev_cat = SEMINAR_CATEGORY_GENITIVE.get(prev_category, "")
        section4.append(
            f"- спортивные судьи по виду спорта «{sport}», имеющие квалификационную категорию спортивного "
            f"судьи «спортивный судья {prev_cat} категории», претендующие на подтверждение квалификационной "
            f"категории «спортивный судья {prev_cat} категории» или присвоение квалификационной категории "
            f"«спортивный судья {cat} категории»;"
        )
    else:
        section4.append(
            f"- спортсмены, специалисты, тренеры, учителя общеобразовательных школ и прочие лица, "
            f"пожелавшие принять участие в проведении соревнований по {sport_dat}, без опыта судейства, "
            f"претендующие на присвоение квалификационной категории «спортивный судья {cat} категории»;"
        )
    section4.append(
        f"- спортивные судьи по виду спорта «{sport}», имеющие квалификационную категорию спортивного "
        f"судьи «спортивный судья {cat} категории», претендующие на подтверждение квалификационной "
        f"категории «спортивный судья {cat} категории»."
    )
    section4.append(
        f"4.3. Все участники семинара для прохождения теоретической подготовки и прохождения "
        f"квалификационного зачета должны иметь возможность подключения к информационному ресурсу "
        f"{fed_short} через сеть Интернет посредством аудио и видео связи и через компьютер для участия "
        f"в практических занятиях."
    )
    section4.append(
        "4.4. Ссылка на информационный ресурс и порядок подключения к нему участников будет выслан "
        "слушателям после подачи заявки на участие в семинаре."
    )

    # 5. Программа семинара
    section5 = []
    if hours_value:
        section5.append(
            f"5.1. Теоретические занятия по подготовке спортивных судей {cat} категории проводится в форме "
            f"семинара и практических занятий по {hours_adj} программе."
        )
    if seminar.qualification_exam == "Да":
        section5.append(
            "5.2. Квалификационный зачет для присвоения или подтверждения квалификационной категории "
            "спортивного судьи проводится в форме экзамена по тестовым вопросам."
        )
    section5.append(
        "5.4. По результатам оценки прохождения семинара участниками, руководителем будет составлен "
        "протокол учета теоретической подготовки и сдачи квалификационного зачета. На основании данного "
        "протокола будут внесены соответствующие записи в карточки учета судейской деятельности спортивного "
        "судьи и в квалификационные книжки спортивных судей."
    )

    # 6. Подведение итогов семинара
    section6 = []
    if hours_value:
        section6.append(
            f"6.1. В программе семинара аттестационная комиссия проводит квалификационный зачёт. Информация "
            f"о теоретической подготовке (лекций и практических занятий в объеме {hours_num} академических "
            f"часов) и о сдаче квалификационного зачета для присвоения (подтверждения) квалификационной "
            f"категории «спортивный судья {cat} категории» в соответствии с квалификационными требованиями "
            f"к спортивным судьям по виду спорта «{sport}», вносится в Протокол учета теоретической "
            f"подготовки и сдачи квалификационного зачета, который рассылается участникам и преподавателям "
            f"семинара."
        )

    # 7. Условия приёма участников
    section7 = []
    if seminar.polozhenie_fee_amount:
        section7.append(
            f"7.1. Все участники семинара обязаны оплатить целевой взнос за участие в семинаре в размере: "
            f"{_clean(seminar.polozhenie_fee_amount)} рублей."
        )
    if seminar.polozhenie_fee_requisites:
        section7.append(f"Реквизиты для оплаты: {_clean(seminar.polozhenie_fee_requisites)}.")
    if seminar.polozhenie_fee_purpose:
        section7.append(f"Назначение платежа: {_clean(seminar.polozhenie_fee_purpose)}.")
    if seminar.polozhenie_accommodation:
        section7.append(f"7.2. Варианты проживания участников: {_clean(seminar.polozhenie_accommodation)}.")
    if seminar.polozhenie_travel:
        section7.append(f"7.3. Проезд до места проведения семинара: {_clean(seminar.polozhenie_travel)}.")

    # 8. Финансирование
    section8 = [
        f"8.1. Расходы, связанные с организацией и проведением семинара, несёт {fed_short}.",
        "8.2. Расходы участников, связанные с участием в семинаре, несут командирующие организации.",
    ]

    # 9. Заявки на участие
    section9 = []
    deadline = _fmt_date_ru(seminar.polozhenie_applications_deadline)
    if deadline:
        section9.append(f"9.1. Для участия в семинаре необходимо до {deadline} включительно:")
    if seminar.polozhenie_applications_email:
        section9.append(f"- отправить заявку на семинар на электронный адрес: {_clean(seminar.polozhenie_applications_email)}.")
    if deadline:
        section9.append("- оплатить целевой взнос за участие в семинаре.")
    if seminar.polozhenie_applications_contacts:
        section9.append(f"9.2. Контакты организаторов: {_clean(seminar.polozhenie_applications_contacts)}.")

    sections = [
        (1, "Цели и задачи", section1),
        (2, "Время и место проведения", section2),
        (3, "Организаторы", section3),
        (4, "Требования к участникам семинара", section4),
        (5, "Программа семинара", section5),
        (6, "Подведение итогов семинара", section6),
        (7, "Условия приёма участников", section7),
        (8, "Финансирование", section8),
        (9, "Заявки на участие", section9),
    ]

    approver_line = " ".join(
        p for p in (seminar.polozhenie_federation_leader_position, _org_name_dative(fed_full)) if p
    )

    return {
        "approver_line": approver_line,
        "approver_signature": _initials_from_full_name(seminar.polozhenie_federation_leader_name),
        "signing_date": _fmt_date_ru(seminar.polozhenie_signing_date),
        "title_sub": title_sub,
        "sections": sections,
        "closing": "Данное положение является приглашением на семинар",
        "lecturers": lecturers,
    }


def _add_paragraph(document, text="", bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line_indent=None, space_after=0):
    p = document.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return p


def _add_table_cell_text(cell, lines):
    cell.paragraphs[0].text = ""
    for extra_p in cell.paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)
    paragraph = cell.paragraphs[0]
    first = True
    for line in lines:
        p = paragraph if first else cell.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def _add_approval_table(document, approver_line, signature, signing_date):
    """Безрамочная таблица 1×3 (47% / 3% / 50%) — гриф «УТВЕРЖДАЮ» в правой ячейке,
    как в образце data/Положение семинар СС1К.doc."""
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    content_width_cm = 21.0 - 2.5 - 1.25
    widths = [Cm(content_width_cm * 0.472), Cm(content_width_cm * 0.028), Cm(content_width_cm * 0.5)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width

    lines = ["УТВЕРЖДАЮ"]
    if approver_line:
        lines.append(approver_line)
    lines.append("")
    lines.append(f"_______________ {signature}".rstrip())
    if signing_date:
        lines.append("")
        lines.append(signing_date)

    _add_table_cell_text(table.rows[0].cells[0], [""])
    _add_table_cell_text(table.rows[0].cells[1], [""])
    _add_table_cell_text(table.rows[0].cells[2], lines)


def generate_polozhenie(seminar, lecturers):
    data = build_polozhenie_data(seminar, lecturers)
    document = docx.Document()

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.25)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    _add_approval_table(document, data["approver_line"], data["approver_signature"], data["signing_date"])
    _add_paragraph(document, "")

    _add_paragraph(document, "ПОЛОЖЕНИЕ", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, data["title_sub"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, "")

    for number, heading, paragraphs in data["sections"]:
        _add_paragraph(document, f"{number}. {heading}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        for line in paragraphs:
            _add_paragraph(document, line, first_line_indent=1.25)
        if not paragraphs:
            _add_paragraph(document, "")
        _add_paragraph(document, "")

    _add_paragraph(document, data["closing"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    if lecturers:
        document.add_page_break()
        _add_paragraph(document, "Приложение 1", align=WD_ALIGN_PARAGRAPH.RIGHT)
        _add_paragraph(document, f"к положению {data['title_sub']}", align=WD_ALIGN_PARAGRAPH.RIGHT)
        _add_paragraph(document, "")
        _add_paragraph(document, "ПРЕПОДАВАТЕЛЬСКИЙ СОСТАВ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = [
            "№ п/п", "Фамилия, имя, отчество", "Городской округ",
            "Квалификационная категория спортивного судьи", "Количество часов",
        ]
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)

        for i, lecturer in enumerate(lecturers, start=1):
            row = table.add_row().cells
            values = [
                str(i), lecturer.full_name or "", lecturer.region or "",
                lecturer.qualification or "", lecturer.lecture_hours or "",
            ]
            for cell, value in zip(row, values):
                cell.text = value
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def polozhenie_filename(seminar):
    name = seminar.name or "семинар"
    return f"Положение {name}.docx"
